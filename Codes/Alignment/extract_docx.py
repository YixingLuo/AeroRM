import re
import json
import os
import shutil
# import pandas as pd
import docx
from docx.document import Document
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
import pandas as pd

json_data = []
chars_to_remove = [' ','.','——','	']

def remove_special_chars(s, chars):
    for char in chars:
        s = s.replace(char,"")
    return s

def get_tabel_dataframe(table):
    data = []
    keys = None
    max_len = -1
    for i, row in enumerate(table.rows):
        text = []
        try: 
            row.cells
        except:
            continue
        for cell in row.cells:
            if not text or str(cell.text) != str(text[-1]):
                text.append(cell.text)
        if len(text) > max_len:
            max_len = len(text)
        data.append(text)
    for text in data:
        if len(text) < max_len:
            text.extend(['' for _ in range(max_len-len(text))])
    df = pd.DataFrame(data)
    csv_data = df.to_csv(sep='|', index=False, header=False)
    return csv_data


def get_heading_numbers(level, level_counters, level_names):
    level_counters[level-1] += 1
    for i in range(level, 6):
        level_counters[i] = 0
    for i in range(level, 6):
        level_names[i]=''
    heading_number = '.'.join(str(level_counters[i]) for i in range(level) if level_counters[i] > 0)
    heading_names = [str(level_names[i]) for i in range(level) if level_names[i] !=''][-1]
    # heading_names = '-'.join(str(level_names[i]) for i in range(level) if level_names[i] !='')
    clean_heading_names = remove_special_chars(heading_names, chars_to_remove)
 
    clean_label = clean_heading_names

    return heading_number, clean_label

def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        if isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_headings_and_contents(file_path):
    doc = docx.Document(file_path)
    headings = []
    content = []
    current_heading = None
    current_content = []
    level_counters = [0]*6
    level_names = ['']*6

    # for para in doc.paragraphs:
    for para in iter_block_items(doc):
        # if (para.style.name.startswith('Heading') or para.style.name.startswith('标题')) and para.text:
        if ('Heading' in para.style.name or '标题' in para.style.name) and para.text:
            # level = int(para.style.name.split(' ')[1])
            try:
                level = int(re.findall(r'(Heading|标题)\s*(\d)', para.style.name)[0][1])
                if level == 1 and len(re.findall(r'目\s*录', para.text)):
                    continue
            except:
                continue
            # print(para.style.name.split(' '), level, para.text, file_path)
            if level > 6:
                continue
            
            level_names[level-1] = para.text
            heading_number, heading_names = get_heading_numbers(level, level_counters, level_names)
            # if level_counters[0] == 1:
            if current_heading:
                headings.append(current_heading)
                content.append('\n'.join(current_content))
            current_heading = {'number':heading_number, 'title': heading_names}
            current_content = []

        elif hasattr(para, 'text') and para.text.strip():
            current_content.append(para.text)
        elif ('Table' in para.style.name) or ('Light List Accent' in para.style.name) or ('MyTab' in para.style.name): # 添加表格
            current_content.append(get_tabel_dataframe(para))

    if current_heading:
        headings.append(current_heading)
        content.append('\n'.join(current_content[1:]))

    return headings, content

def process_document(file_path, hastitle=False, recursion=False):
    
    headings, contents = get_headings_and_contents(file_path)
    data = []

    if hastitle:
        for i in range(len(headings)):
            contents[i] = headings[i]['title'] + '\n' + contents[i]

    if recursion:
        for i in range(len(headings)):
            for j in range(i+1, len(headings)):
                if headings[i]['number'] in headings[j]['number']:
                    contents[i] += '\n' + contents[j]
                else:
                    break

    for heading, content in zip(headings, contents):
        data.append({
                    'number': heading['number'],
                    'title': heading['title'],
                    'content': content
        })

    return data

def extract_table(file_path):
    try:
        doc = docx.Document(file_path)
    except:
        return None
    table_list = []

    # for para in doc.paragraphs:
    for para in iter_block_items(doc):
        try:
            if ('Table' in para.style.name) or ('Light List Accent' in para.style.name) or ('MyTab' in para.style.name):
                table_list.append(get_tabel_dataframe(para))
        except:
            continue
    return table_list

def test():
    file_path = r'xxx'

    data = process_document(file_path)
